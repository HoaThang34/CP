from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_document():
    doc = Document()

    # Title
    title = doc.add_heading('GIẢI CHI TIẾT ĐỀ THI HSG BÌNH ĐỊNH', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- DAYBIAN ---
    doc.add_heading('1. Bài DAYBIAN', level=1)
    
    doc.add_heading('Đề bài', level=2)
    p = doc.add_paragraph()
    p.add_run('Cho dãy số A gồm n phần tử và một số nguyên k. Yêu cầu đếm số lượng dãy con liên tiếp có tổng bằng k.\n')
    p.add_run('Input: Dòng đầu n, k. Dòng sau là dãy A.\n')
    p.add_run('Output: Số lượng dãy con có tổng bằng k.')

    doc.add_heading('Hướng dẫn giải', level=2)
    doc.add_paragraph('Bài toán yêu cầu đếm số lượng cặp chỉ số (i, j) sao cho tổng từ A[i] đến A[j] bằng k.')
    
    doc.add_heading('Subtask 1 (N nhỏ)', level=3)
    doc.add_paragraph('Duyệt qua tất cả các cặp (i, j) và tính tổng. Nếu tổng bằng k thì tăng biến đếm. Độ phức tạp O(N^2).')

    doc.add_heading('Full Solution', level=3)
    doc.add_paragraph('Sử dụng kỹ thuật Mảng cộng dồn (Prefix Sum). Gọi S[i] là tổng các phần tử từ A[0] đến A[i].')
    doc.add_paragraph('Tổng đoạn con từ j đến i là S[i] - S[j-1]. Ta cần tìm số lượng j sao cho S[i] - S[j-1] = k <=> S[j-1] = S[i] - k.')
    doc.add_paragraph('Sử dụng Map hoặc mảng đếm để lưu tần suất xuất hiện của các giá trị S[i] khi duyệt. Độ phức tạp O(N log N) hoặc O(N).')

    doc.add_heading('Code tham khảo', level=2)
    code_daybian = """
#include <iostream>
#include <vector>
#include <map>
using namespace std;

// Subtask 1: Brute force O(N^2)
namespace Subtask1 {
    void solve(int n, long long k, const vector<long long>& a) {
        long long count = 0;
        for (int i = 0; i < n; i++) {
            long long current_sum = 0;
            for (int j = i; j < n; j++) {
                current_sum += a[j];
                if (current_sum == k) count++;
            }
        }
        cout << count << endl;
    }
}

// Full: Prefix Sum + Map O(N log N)
namespace Full {
    void solve(int n, long long k, const vector<long long>& a) {
        map<long long, int> prefix_counts;
        prefix_counts[0] = 1; 
        long long current_sum = 0;
        long long ans = 0;
        for (long long x : a) {
            current_sum += x;
            if (prefix_counts.count(current_sum - k)) {
                ans += prefix_counts[current_sum - k];
            }
            prefix_counts[current_sum]++;
        }
        cout << ans << endl;
    }
}

int main() {
    ios_base::sync_with_stdio(false); cin.tie(NULL);
    int n; long long k;
    if (cin >> n >> k) {
        vector<long long> a(n);
        for (int i = 0; i < n; i++) cin >> a[i];
        Full::solve(n, k, a);
    }
    return 0;
}
"""
    p = doc.add_paragraph(code_daybian)
    for run in p.runs:
        run.font.name = 'Courier New'

    # --- HOPQUA ---
    doc.add_page_break()
    doc.add_heading('2. Bài HOPQUA', level=1)
    
    doc.add_heading('Đề bài', level=2)
    p = doc.add_paragraph()
    p.add_run('Có n hộp quà kích thước A và m yêu cầu kích thước tối thiểu B. Mỗi hộp chỉ dùng 1 lần. Chọn hộp sao cho tổng kích thước các hộp được chọn là nhỏ nhất.\n')
    p.add_run('Input: n, m. Mảng A, Mảng B.\n')
    p.add_run('Output: Tổng kích thước nhỏ nhất hoặc -1.')

    doc.add_heading('Hướng dẫn giải', level=2)
    doc.add_paragraph('Bài toán sử dụng giải thuật Tham lam (Greedy).')

    doc.add_heading('Full Solution', level=3)
    doc.add_paragraph('Sắp xếp cả hai mảng A (hộp quà) và B (yêu cầu) theo thứ tự tăng dần.')
    doc.add_paragraph('Với mỗi yêu cầu B[i], ta tìm hộp quà A[j] nhỏ nhất sao cho A[j] >= B[i].')
    doc.add_paragraph('Sử dụng kỹ thuật 2 con trỏ để duyệt hiệu quả. Độ phức tạp O(N log N + M log M).')

    doc.add_heading('Code tham khảo', level=2)
    code_hopqua = """
#include <iostream>
#include <vector>
#include <algorithm>
using namespace std;

void solve(int n, int m, vector<long long>& A, vector<long long>& B) {
    sort(A.begin(), A.end());
    sort(B.begin(), B.end());

    int ptr_a = 0, matches = 0;
    long long total_cost = 0;

    for (int i = 0; i < m; i++) {
        while (ptr_a < n && A[ptr_a] < B[i]) {
            ptr_a++;
        }
        if (ptr_a < n) {
            total_cost += A[ptr_a];
            matches++;
            ptr_a++;
        } else {
            cout << -1 << endl;
            return;
        }
    }
    cout << total_cost << endl;
}

int main() {
    ios_base::sync_with_stdio(false); cin.tie(NULL);
    int n, m;
    if (cin >> n >> m) {
        vector<long long> A(n), B(m);
        for (int i = 0; i < n; i++) cin >> A[i];
        for (int i = 0; i < m; i++) cin >> B[i];
        solve(n, m, A, B);
    }
    return 0;
}
"""
    p = doc.add_paragraph(code_hopqua)
    for run in p.runs:
        run.font.name = 'Courier New'

    # --- MATTHU ---
    doc.add_page_break()
    doc.add_heading('3. Bài MATTHU', level=1)
    
    doc.add_heading('Đề bài', level=2)
    p = doc.add_paragraph()
    p.add_run('Thực hiện 3 loại truy vấn trên dãy A: (1) Tính tổng đoạn [l, r], (2) Gán A[i] = A[i] % w trong đoạn [l, r], (3) Gán A[x] = k.\n')
    
    doc.add_heading('Hướng dẫn giải', level=2)
    
    doc.add_heading('Subtask 1 (N, M <= 1000)', level=3)
    doc.add_paragraph('Duyệt trâu cho mỗi truy vấn. Cập nhật trực tiếp trên mảng và tính tổng bằng vòng lặp. Độ phức tạp O(M*N).')

    doc.add_heading('Full Solution', level=3)
    doc.add_paragraph('Sử dụng Segment Tree. Ngoài lưu tổng (sum), mỗi node lưu thêm giá trị lớn nhất (max) trong đoạn.')
    doc.add_paragraph('Khi cập nhật modulo w: Nếu max của đoạn < w, phép modulo không thay đổi giá trị nào -> bỏ qua đệ quy xuống node con. Điều này giúp giảm độ phức tạp vì số lần một số giảm giá trị qua phép modulo là logarit.')

    doc.add_heading('Code tham khảo', level=2)
    code_matthu = """
// Xem file Full Solution trong code mẫu
// Key function: update_modulo
void update_modulo(int node, int start, int end, int l, int r, long long k) {
    if (start > end || start > r || end < l || tree_max[node] < k) return;
    if (start == end) {
        tree_sum[node] %= k;
        tree_max[node] %= k;
        return;
    }
    // ... recurse children ...
    tree_sum[node] = tree_sum[2*node] + tree_sum[2*node+1];
    tree_max[node] = max(tree_max[2*node], tree_max[2*node+1]);
}
"""
    p = doc.add_paragraph(code_matthu)
    for run in p.runs:
        run.font.name = 'Courier New'

    # --- THUTHACH ---
    doc.add_page_break()
    doc.add_heading('4. Bài THUTHACH', level=1)
    
    doc.add_heading('Đề bài', level=2)
    p = doc.add_paragraph()
    p.add_run('Tìm số nguyên dương thứ k không chia hết cho a và không chia hết cho b.\n')
    
    doc.add_heading('Hướng dẫn giải', level=2)
    
    doc.add_heading('Subtask 1 (k, a, b <= 100)', level=3)
    doc.add_paragraph('Duyệt số x từ 1 tăng dần, kiểm tra điều kiện chia hết, đếm số lượng thoả mãn cho đến khi đủ k số.')

    doc.add_heading('Full Solution', level=3)
    doc.add_paragraph('Tìm kiếm nhị phân kết quả (Binary Search on Answer).')
    doc.add_paragraph('Hàm check(X): đếm số lượng số <= X thoả mãn điều kiện. Sử dụng nguyên lý bù trừ:')
    doc.add_paragraph('Count = X - (X/a + X/b - X/LCM(a,b))')
    doc.add_paragraph('Nếu Count >= k thì giảm cận trên, ngược lại tăng cận dưới. Độ phức tạp O(log(MAX)).')

    doc.add_heading('Code tham khảo', level=2)
    code_thuthach = """
long long lcm(long long a, long long b) { return (a / gcd(a, b)) * b; }

long long count_valid(long long x, long long a, long long b, long long l_val) {
    long long div_any = x/a + x/b - x/l_val;
    return x - div_any;
}

void solve(long long k, long long a, long long b) {
    long long l_val = lcm(a, b);
    long long low = 1, high = 4e18, ans = -1;
    while (low <= high) {
        long long mid = low + (high - low) / 2;
        if (count_valid(mid, a, b, l_val) >= k) {
            ans = mid; high = mid - 1;
        } else low = mid + 1;
    }
    cout << ans << endl;
}
"""
    p = doc.add_paragraph(code_thuthach)
    for run in p.runs:
        run.font.name = 'Courier New'

    doc.save('GIAI_CHI_TIET_HSG_BINH_DINH.docx')
    print("Document created successfully.")

if __name__ == "__main__":
    create_document()
