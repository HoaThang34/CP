from docx import Document
from docx.shared import Pt
import os

problem_name = "KHOANGCACH"
title = "HƯỚNG DẪN GIẢI BÀI: KHOẢNG CÁCH"

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
CODE_PATH = os.path.join(SCRIPT_DIR, "CODE_AC", f"{problem_name}.cpp")
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "huongdangiai.docx")

def main():
    doc = Document()
    doc.add_heading(title, 0)

    doc.add_heading('1. Xác định bài toán', level=1)
    doc.add_paragraph('Cho xâu S gồm các kí tự in thường (a-z). Có Q truy vấn, mỗi truy vấn cho đoạn [L, R].')
    doc.add_paragraph('Khoảng cách giữa 2 ký tự c1, c2 trên vòng tròn bảng chữ cái là min(|c1-c2|, 26 - |c1-c2|).')
    doc.add_paragraph('Khoảng cách của một xâu con là khoảng cách lớn nhất giữa hai ký tự bất kỳ có mặt trong xâu con đó.')
    doc.add_paragraph('Yêu cầu: Tính khoảng cách cho từng truy vấn.')
    
    doc.add_heading('Input', level=2)
    doc.add_paragraph('- Dòng 1: Xâu S.')
    doc.add_paragraph('- Dòng 2: Q.')
    doc.add_paragraph('- Q dòng tiếp theo: L, R.')
    
    doc.add_heading('Output', level=2)
    doc.add_paragraph('Q dòng kết quả.')

    doc.add_heading('2. Thuật toán', level=1)
    doc.add_paragraph('Do chỉ có 26 ký tự (a-z), ta có thể sử dụng Mảng cộng dồn (Prefix Sum) để đếm số lần xuất hiện của từng ký tự.')
    doc.add_paragraph('Gọi Count[c][i] là số lần ký tự c xuất hiện trong S[0...i].')
    doc.add_paragraph('Trong đoạn [L, R], ký tự c xuất hiện nếu Count[c][R] - Count[c][L-1] > 0.')
    doc.add_paragraph('Với mỗi truy vấn:')
    doc.add_paragraph('  Bước 1: Lấy danh sách các ký tự xuất hiện trong đoạn [L, R].')
    doc.add_paragraph('  Bước 2: Nếu số lượng ký tự < 2, kết quả là 0.')
    doc.add_paragraph('  Bước 3: Nếu >= 2, duyệt tất cả các cặp ký tự có mặt và tìm max khoảng cách.')
    doc.add_paragraph('Độ phức tạp mỗi truy vấn là O(26^2), rất nhỏ.')

    doc.add_heading('3. Đánh giá', level=1)
    doc.add_paragraph('Tiền xử lý: O(26 * N).')
    doc.add_paragraph('Truy vấn: O(Q * 26^2).')
    doc.add_paragraph('Tổng quát: O(N + Q). Chạy tốt với N, Q = 10^5.')

    doc.add_heading('4. Code mẫu C++', level=1)
    
    code = ""
    if os.path.exists(CODE_PATH):
        with open(CODE_PATH, 'r') as f:
            code = f.read()
            
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    
    doc.save(OUTPUT_PATH)
    print(f"Docs saved to {OUTPUT_PATH}")

if __name__ == "__main__":
    main()
