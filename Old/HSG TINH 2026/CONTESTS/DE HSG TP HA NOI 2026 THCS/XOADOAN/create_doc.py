from docx import Document
from docx.shared import Pt
import os

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
CODE_PATH = os.path.join(SCRIPT_DIR, "CODE_AC", "XOADOAN.cpp")
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "huongdangiai.docx")

def main():
    doc = Document()
    
    # Title
    heading = doc.add_heading('HƯỚNG DẪN GIẢI BÀI TẬP: XÓA ĐOẠN', level=1)
    
    # 1. Xác định bài toán
    doc.add_heading('1. Xác định bài toán', level=2)
    doc.add_paragraph('Cho dãy số nguyên A gồm N phần tử và số nguyên S.')
    doc.add_paragraph('Yêu cầu: Tìm độ dài nhỏ nhất của đoạn con liên tiếp cần xóa sao cho tổng các phần tử còn lại không vượt quá S. Nếu không có cách xóa, in -1.')
    
    doc.add_heading('Input:', level=3)
    doc.add_paragraph('- Dòng 1: N, S')
    doc.add_paragraph('- Dòng 2: Dãy A')
    
    doc.add_heading('Output:', level=3)
    doc.add_paragraph('- Đô dài nhỏ nhất của đoạn con cần xóa.')
    
    doc.add_heading('Ràng buộc:', level=3)
    doc.add_paragraph('- N <= 10^5')
    doc.add_paragraph('- |Ai| <= 10^9')
    doc.add_paragraph('- |S| <= 10^14')

    # 2. Thuật toán
    doc.add_heading('2. Thuật toán', level=2)
    doc.add_paragraph('Bài toán tương đương với việc tìm đoạn con có tổng lớn nhất mà tổng đó >= Tổng dãy - S.')
    doc.add_paragraph('Gọi K = Tổng dãy - S. Ta cần tìm đoạn con ngắn nhất có tổng >= K.')
    
    doc.add_paragraph('Các bước thực hiện:')
    doc.add_paragraph('Bước 1: Tính tổng toàn bộ dãy số (TotalSum). Tính K = TotalSum - S.')
    doc.add_paragraph('Bước 2: Nếu TotalSum <= S, kết quả là 0 (không cần xóa gì).')
    doc.add_paragraph('Bước 3: Xây dựng mảng cộng dồn P, với P[i] là tổng các phần tử từ A[0] đến A[i-1].')
    doc.add_paragraph('Bước 4: Sử dụng cấu trúc hàng đợi hai đầu (Deque) để tìm đoạn con ngắn nhất [i, j) sao cho P[j] - P[i] >= K.')
    doc.add_paragraph('   - Duyệt j từ 0 đến N.')
    doc.add_paragraph('   - Kiểm tra đầu hàng đợi: Nếu P[j] - P[dq.front()] >= K, cập nhật kết quả min_len = min(min_len, j - dq.front()) và loại bỏ dq.front().')
    doc.add_paragraph('   - Duyệt trì tính đơn điệu: Loại bỏ các phần tử cuối hàng đợi nếu P[dq.back()] >= P[j] (vì j tốt hơn về vị trí và cả về giá trị P).')
    doc.add_paragraph('   - Thêm j vào cuối hàng đợi.')
    doc.add_paragraph('Bước 5: Nếu không tìm thấy đoạn nào thỏa mãn và không xóa hết được (trường hợp S < 0), in -1.')

    # 3. Đánh giá thuật toán
    doc.add_heading('3. Đánh giá thuật toán', level=2)
    doc.add_paragraph('Độ phức tạp thời gian: O(N). Do mỗi phần tử chỉ được chèn và lấy ra khỏi Deque tối đa 1 lần.')
    doc.add_paragraph('Độ phức tạp bộ nhớ: O(N) để lưu mảng cộng dồn và Deque.')
    doc.add_paragraph('Phù hợp với giới hạn N = 10^5 (thời gian chạy ~0.1s). Kiểu dữ liệu long long (64-bit) đủ để chứa tổng lên tới 10^14.')

    # 4. Code mẫu
    doc.add_heading('4. Code mẫu C++', level=2)
    
    code_content = ""
    if os.path.exists(CODE_PATH):
        with open(CODE_PATH, "r", encoding="utf-8") as f:
            code_content = f.read()
    else:
        code_content = "// Code file not found"
        
    p = doc.add_paragraph()
    run = p.add_run(code_content)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

    doc.save(OUTPUT_PATH)
    print(f"Document saved to {OUTPUT_PATH}")

if __name__ == "__main__":
    main()
