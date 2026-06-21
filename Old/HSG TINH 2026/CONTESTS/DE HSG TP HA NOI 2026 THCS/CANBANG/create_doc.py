from docx import Document
from docx.shared import Pt
import os

problem_name = "CANBANG"
title = "HƯỚNG DẪN GIẢI BÀI: CÂN BẰNG"

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
CODE_PATH = os.path.join(SCRIPT_DIR, "CODE_AC", f"{problem_name}.cpp")
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "huongdangiai.docx")

def main():
    doc = Document()
    doc.add_heading(title, 0)

    doc.add_heading('1. Xác định bài toán', level=1)
    doc.add_paragraph('Cho dãy số nguyên phân biệt A gồm N phần tử và số K.')
    doc.add_paragraph('Yêu cầu: Đếm số lượng phần tử A[i] thỏa mãn tính chất: trong dãy tồn tại phần tử (A[i] - K) VÀ tồn tại phần tử (A[i] + K).')
    
    doc.add_heading('Input', level=2)
    doc.add_paragraph('Dòng 1: N và K.')
    doc.add_paragraph('Dòng 2: Dãy A gồm N phần tử phân biệt.')
    
    doc.add_heading('Output', level=2)
    doc.add_paragraph('Số lượng phần tử "cân bằng".')

    doc.add_heading('2. Thuật toán', level=1)
    doc.add_paragraph('Bước 1: Đọc dãy số và lưu tất cả các phần tử vào một cấu trúc dữ liệu cho phép tra cứu nhanh (Hash Set hoặc Balanced Tree Set). Trong C++, dùng std::set hoặc std::unordered_set.')
    doc.add_paragraph('Bước 2: Duyệt qua từng phần tử x trong dãy ban đầu.')
    doc.add_paragraph('Bước 3: Với mỗi x, kiểm tra xem (x - K) và (x + K) có tồn tại trong tập hợp hay không.')
    doc.add_paragraph('Bước 4: Nếu cả hai đều tồn tại, tăng biến đếm.')
    doc.add_paragraph('Bước 5: In kết quả.')

    doc.add_heading('3. Đánh giá', level=1)
    doc.add_paragraph('Độ phức tạp thời gian: O(N) nếu dùng Hash Set, hoặc O(N log N) nếu dùng Tree Set.')
    doc.add_paragraph('Độ phức tạp bộ nhớ: O(N) để lưu tập hợp.')
    doc.add_paragraph('Giải thuật tối ưu đảm bảo chạy tốt với N=10^5 trong thời gian giới hạn.')

    doc.add_heading('4. Code mẫu C++', level=1)
    
    code = ""
    if os.path.exists(CODE_PATH):
        with open(CODE_PATH, 'r') as f:
            code = f.read()
            
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    
    doc.save(OUTPUT_PATH)
    print(f"Docs saved to {OUTPUT_PATH}")

if __name__ == "__main__":
    main()
