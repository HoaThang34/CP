from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUTPUT_FILENAME = "HUONGDANGIAI_CHITIET_FULL.docx"
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))

def add_problem(doc, problem_data):
    # Title
    p = doc.add_heading(problem_data['title'], level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. Problem Statement
    doc.add_heading('1. Xác định bài toán', level=2)
    for line in problem_data['statement']:
        doc.add_paragraph(line)
        
    doc.add_heading('Input:', level=3)
    for line in problem_data['input']:
        doc.add_paragraph(line, style='List Bullet')
        
    doc.add_heading('Output:', level=3)
    for line in problem_data['output']:
        doc.add_paragraph(line, style='List Bullet')

    # 2. Algorithm
    doc.add_heading('2. Thuật toán', level=2)
    for line in problem_data['algo']:
        doc.add_paragraph(line)

    # 3. Complexity
    doc.add_heading('3. Đánh giá thuật toán', level=2)
    for line in problem_data['complexity']:
        doc.add_paragraph(line)

    # 4. Code
    doc.add_heading('4. Code mẫu C++', level=2)
    
    code_path = os.path.join(SCRIPT_DIR, problem_data['folder'], "CODE_AC", f"{problem_data['code_file']}")
    code_content = "// Code not found"
    if os.path.exists(code_path):
        with open(code_path, "r", encoding="utf-8") as f:
            code_content = f.read()
            
    p = doc.add_paragraph()
    run = p.add_run(code_content)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    # Generic syntax highlighting sim not possible easily in docx without complex logic, keep plain mono.

    doc.add_page_break()

def main():
    doc = Document()
    
    # Main Title
    title = doc.add_heading('HƯỚNG DẪN GIẢI ĐỀ THI HSG TP HÀ NỘI 2026', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # Data
    problems = [
        {
            'folder': 'CHOXUAN',
            'code_file': 'CHOXUAN.cpp',
            'title': 'BÀI I: CHỢ XUÂN',
            'statement': [
                'Lớp An có N đồng. Thuê gian hàng 7 ngày, giá K đồng/ngày.',
                'Yêu cầu: Tính số tiền còn lại sau khi thuê. Nếu không đủ tiền (N < 7*K) thì in -1.'
            ],
            'input': [
                'Dòng 1: Số nguyên dương N (tiền hiện có).',
                'Dòng 2: Số nguyên dương K (giá thuê 1 ngày).'
            ],
            'output': [
                'Một số nguyên duy nhất là tiền còn lại hoặc -1.'
            ],
            'algo': [
                'Bài toán yêu cầu tính toán cơ bản.',
                'Bước 1: Tính tổng chi phí Cost = K * 7.',
                'Bước 2: Kiểm tra nếu N >= Cost thì kết quả là N - Cost.',
                'Bước 3: Ngược lại, kêt quả là -1.',
                'Chú ý: N và K có thể lớn, cần dùng kiểu dữ liệu số nguyên 64-bit (long long).'
            ],
            'complexity': [
                'Thời gian: O(1).',
                'Bộ nhớ: O(1).'
            ]
        },
        {
            'folder': 'CANBANG',
            'code_file': 'CANBANG.cpp',
            'title': 'BÀI II: CÂN BẰNG',
            'statement': [
                'Cho dãy số A gồm N phần tử phân biệt và số nguyên K.',
                'Một phần tử A[i] gọi là "cân bằng K" nếu trong dãy tồn tại cả (A[i] - K) và (A[i] + K).',
                'Yêu cầu: Đếm số lượng phần tử cân bằng.'
            ],
            'input': [
                'N và K.',
                'Dãy số A.'
            ],
            'output': [
                'Số lượng phần tử thỏa mãn.'
            ],
            'algo': [
                'Để kiểm tra sự tồn tại của một giá trị trong dãy nhanh chóng, ta dùng cấu trúc dữ liệu Set (Tập hợp) hoặc Map.',
                'Bước 1: Đọc dãy A và chèn tất cả phần tử vào một Hash Set (std::unordered_set) hoặc Tree Set (std::set).',
                'Bước 2: Khởi tạo biến đếm count = 0.',
                'Bước 3: Duyệt lại từng phần tử x trong dãy A:',
                '   - Kiểm tra nếu (x - K) có trong Set VÀ (x + K) có trong Set.',
                '   - Nếu đúng, tăng count.',
                'Bước 4: In kết quả.'
            ],
            'complexity': [
                'Thời gian: O(N) trung bình nếu dùng Hash Set, hoặc O(N log N) nếu dùng Tree Set.',
                'Bộ nhớ: O(N) để lưu Set.'
            ]
        },
        {
            'folder': 'KHOANGCACH',
            'code_file': 'KHOANGCACH.cpp',
            'title': 'BÀI III: KHOẢNG CÁCH',
            'statement': [
                'Cho xâu S chỉ gồm ký tự thường. Có Q truy vấn đoạn [L, R].',
                'Quy ước các ký tự được xếp trên vòng tròn (a-z). Khoảng cách giữa 2 ký tự la ngắn nhất trên vòng tròn.',
                'Khoảng cách của xâu con là khoảng cách lớn nhất giữa 2 ký tự bất kỳ có mặt trong xâu con đó.',
                'Yêu cầu: Trả lời Q truy vấn.'
            ],
            'input': [
                'Xâu S.',
                'Số truy vấn Q.',
                'Q dòng chứa L, R.'
            ],
            'output': [
                'Q dòng kết quả.'
            ],
            'algo': [
                'Do bảng chữ cái chỉ có 26 ký tự, ta có thể dùng Mảng cộng dồn (Prefix Count) để xác định sự xuất hiện của ký tự trong O(1).',
                'Bước 1: Xây dựng mảng P[char][i] = số lần xuất hiện của char trong S[0..i].',
                'Bước 2: Với mỗi truy vấn [L, R]:',
                '   - Tìm tập hợp các ký tự xuất hiện (nếu P[c][R] - P[c][L-1] > 0).',
                '   - Nếu số ký tự < 2, khoảng cách = 0.',
                '   - Nếu >= 2, duyệt tất cả cặp (c1, c2) trong tập hợp đó và tìm max dist(c1, c2).',
                '   - Vì tối đa chỉ có 26 ký tự, số cặp tối đa là 26*25/2, rất nhỏ.',
            ],
            'complexity': [
                'Tiền xử lý: O(26 * N).',
                'Truy vấn: O(Q * 26^2).',
                'Tổng: O(N + Q) với hằng số nhỏ.'
            ]
        },
        {
            'folder': 'XOADOAN',
            'code_file': 'XOADOAN.cpp',
            'title': 'BÀI IV: XÓA ĐOẠN',
            'statement': [
                'Cho dãy A gồm N phần tử và số S. Cần xóa một đoạn con liên tiếp sao cho tổng các phần tử còn lại <= S.',
                'Yêu cầu: Tìm độ dài đoạn con cần xóa NHỎ NHẤT.'
            ],
            'input': [
                'N và S.',
                'Dãy A.'
            ],
            'output': [
                'Độ dài nhỏ nhất hoặc -1.'
            ],
            'algo': [
                'Bài toán tương đương: Tìm đoạn con [i, j] có tổng lớn nhất sao cho: Tổng(Toàn_bộ) - Tổng(i...j) <= S.',
                'Hay: Tổng(i...j) >= Tổng(Toàn_bộ) - S = K.',
                'Ta cần tìm đoạn con ngắn nhất có tổng >= K.',
                'Dùng kỹ thuật Deque (Monoqueue) trên mảng cộng dồn P:',
                'Bước 1: Tính K. Nếu K <= 0 (tức tổng ban đầu <= S), xóa 0 phần tử.',
                'Bước 2: Duyệt j từ 0 đến N (chỉ số mảng cộng dồn).',
                '   - Trong Deque lưu các chỉ số i sao cho P[i] tăng dần.',
                '   - Tại j, nếu P[j] - P[dq.front()] >= K, cập nhật min_len = j - dq.front() và pop front (đã tìm thấy đoạn ngắn nhất bắt đầu từ front).',
                '   - Duy trì tính tăng dần của P trong Deque khi push j.',
            ],
            'complexity': [
                'Thời gian: O(N). Mỗi phần tử vào/ra Deque 1 lần.',
                'Bộ nhớ: O(N).'
            ]
        },
        {
            'folder': 'BANSUNG',
            'code_file': 'BANSUNG.cpp',
            'title': 'BÀI V: BẮN SÚNG',
            'statement': [
                'Có N bia với độ bên A[i]. Cần tìm sức công phá X nhỏ nhất để phá hủy hết bia trong <= K lần bắn.',
                'Cơ chế: Bắn bia đầu tiên còn sống. Sát thương lan ra sau: max(0, X - (kc)^2).',
            ],
            'input': [
                'N, K.',
                'Dãy A.'
            ],
            'output': [
                'Min X.'
            ],
            'algo': [
                'Nhận thấy tính chất đơn điệu: Nếu X phá được thì X+1 cũng phá được.',
                '-> Sử dụng Tìm kiếm nhị phân kết quả (Binary Search on Answer).',
                'Phạm vi X: [1, 2*10^14].',
                'Hàm check(Val):',
                '   - Mô phỏng quá trình bắn với X = Val.',
                '   - Dùng vòng lặp bắn tối đa K lần.',
                '   - Mỗi lần tìm bia đầu tiên dương, trừ máu các bia sau bị ảnh hưởng.',
                '   - Nếu sau K lần vẫn còn bia dương -> False.',
                '   - Tối ưu mô phỏng: Vòng lặp trừ máu chỉ chạy đến khi sát thương = 0 (cự ly sqrt(X)).',
            ],
            'complexity': [
                'Thời gian: O(N * K * log(MaxX)).',
                'Với N=2e5, K=10, thuật toán đủ nhanh.'
            ]
        }
    ]

    for p in problems:
        add_problem(doc, p)
        print(f"Added {p['folder']}")

    doc.save(OUTPUT_FILENAME)
    print(f"Full documentation saved to {OUTPUT_FILENAME}")

if __name__ == "__main__":
    main()
