from docx import Document
from docx.shared import Pt
import os

problem_name = "BANSUNG"
title = "HƯỚNG DẪN GIẢI BÀI: BẮN SÚNG"

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
CODE_PATH = os.path.join(SCRIPT_DIR, "CODE_AC", f"{problem_name}.cpp")
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "huongdangiai.docx")

def main():
    doc = Document()
    doc.add_heading(title, 0)

    doc.add_heading('1. Xác định bài toán', level=1)
    doc.add_paragraph('Cho N tấm bia với độ bền A[i]. Xạ thủ chọn sức công phá X. Mỗi lần bắn vào bia đầu tiên còn "sống" (A[i] > 0).')
    doc.add_paragraph('Viên đạn gây sát thương lan ra các bia phía sau: bia thứ j (j >= i) chịu sát thương max(0, X - (j-i)^2).')
    doc.add_paragraph('Yêu cầu: Tìm X nhỏ nhất để phá hủy toàn bộ N tấm bia trong tối đa K lần bắn.')
    
    doc.add_heading('Input', level=2)
    doc.add_paragraph('- Dòng 1: N, K.')
    doc.add_paragraph('- Dòng 2: Dãy A.')
    
    doc.add_heading('Output', level=2)
    doc.add_paragraph('Giá trị X nhỏ nhất.')

    doc.add_heading('2. Thuật toán', level=1)
    doc.add_paragraph('Nhận xét: Nếu X đủ lớn để phá hủy trong K lần, thì X+1 cũng thỏa mãn. Hàm tính chất đơn điệu -> Dùng Tìm kiếm nhị phân theo kết quả (Binary Search on Answer).')
    doc.add_paragraph('Phạm vi tìm kiếm X: từ 1 đến giá trị đủ lớn (ví dụ 10^14).')
    doc.add_paragraph('Hàm kiểm tra check(X):')
    doc.add_paragraph('  - Mô phỏng quá trình bắn.')
    doc.add_paragraph('  - Lặp K lần bắn (hoặc đến khi hết bia).')
    doc.add_paragraph('  - Mỗi lần bắn, tìm bia đầu tiên chưa bị phá hủy. Trừ độ bền các bia bị ảnh hưởng.')
    doc.add_paragraph('  - Lưu ý tối ưu vòng lặp trừ độ bền: chỉ cần duyệt đến khi X - (j-i)^2 <= 0 (tức là j - i > sqrt(X)).')
    doc.add_paragraph('  - Sau K lần, nếu tất cả bia có A[i] <= 0 thì trả về True.')
    
    doc.add_heading('3. Đánh giá', level=1)
    doc.add_paragraph('Độ phức tạp: O(N * K * log(MaxX)).')
    doc.add_paragraph('Với N=2*10^5, K=10, log(MaxX) ≈ 60: Tổng phép tính khoảng 10^8, chấp nhận được trong thời gian 1s.')

    doc.add_heading('4. Code mẫu C++', level=1)
    
    code = ""
    if os.path.exists(CODE_PATH):
        with open(CODE_PATH, 'r') as f:
            code = f.read()
            
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    
    doc.save(OUTPUT_PATH)
    print(f"Docs saved to {OUTPUT_PATH}")

if __name__ == "__main__":
    main()
