from docx import Document
from docx.shared import Pt
import os

problem_name = "CHOXUAN"
title = "HƯỚNG DẪN GIẢI BÀI: CHỢ XUÂN"

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
CODE_PATH = os.path.join(SCRIPT_DIR, "CODE_AC", f"{problem_name}.cpp")
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "huongdangiai.docx")

def main():
    doc = Document()
    doc.add_heading(title, 0)

    doc.add_heading('1. Xác định bài toán', level=1)
    doc.add_paragraph('Lớp An có N đồng. Thuê gian hàng 7 ngày, mỗi ngày giá K đồng.')
    doc.add_paragraph('Yêu cầu: Tính số tiền còn lại. Nếu không đủ tiền thuê 7 ngày thì in -1.')
    
    doc.add_heading('Input', level=2)
    doc.add_paragraph('Hai dòng chứa N và K.')
    
    doc.add_heading('Output', level=2)
    doc.add_paragraph('Số tiền còn lại hoặc -1.')

    doc.add_heading('2. Thuật toán', level=1)
    doc.add_paragraph('Bước 1: Tính tổng chi phí cần thiết: TotalCost = K * 7.')
    doc.add_paragraph('Bước 2: So sánh số tiền hiện có N với TotalCost.')
    doc.add_paragraph('Bước 3: Nếu N >= TotalCost, in ra N - TotalCost. Ngược lại in ra -1.')

    doc.add_heading('3. Đánh giá', level=1)
    doc.add_paragraph('Độ phức tạp thời gian: O(1).')
    doc.add_paragraph('Độ phức tạp bộ nhớ: O(1).')
    doc.add_paragraph('Ràng buộc: N, K lớn nên cần dùng kiểu dữ liệu số nguyên 64-bit (long long trong C++).')

    doc.add_heading('4. Code mẫu C++', level=1)
    
    code = ""
    if os.path.exists(CODE_PATH):
        with open(CODE_PATH, 'r') as f:
            code = f.read()
            
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    
    doc.save(OUTPUT_PATH)
    print(f"Docs saved to {OUTPUT_PATH}")

if __name__ == "__main__":
    main()
